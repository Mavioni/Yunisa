from docx import Document
import sys

def read_docx(path, out_path):
    doc = Document(path)
    with open(out_path, 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            text = para.text.strip().replace('\r', '')
            if text:
                f.write(text + '\n')
        # Also extract tables just in case
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip().replace('\r', ''))
                f.write(' | '.join(row_data) + '\n')

if __name__ == "__main__":
    if len(sys.argv) > 2:
        read_docx(sys.argv[1], sys.argv[2])
